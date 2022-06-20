from app.feedback.retrieveText.convertDocxTxtToText import getDOCXText, subtractTextFromParagraph
import os
import docx


def testGetDocxHeading(testClient):
    """
        Test if the getDOCXText function retrieves the text from the docx file without headings.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'headingTest.docx')

    text, references = getDOCXText(fileDir)
    assert text == 'This is some text.\n\nMore text.'
    assert references == ''


def testGetDocxReferences(testClient):
    """
        Test if the getDOCXText function retrieves the text from the docx file without references.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'referencesTest.docx')

    text, references = getDOCXText(fileDir)
    assert text == 'Text.\n\nMore text.\n\nNew text.'
    assert references == 'Reference 1\n\nReference 2\n\nReference 1\n\nReference 2'


def testGetDocxImages(testClient):
    """
        Test if the getDOCXText function retrieves the text from the docx file without images and corresponding captions.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'imagesTest.docx')

    text, references = getDOCXText(fileDir)
    assert text == 'Lorem ipsum dolor sit amet, consectetuer adipiscing elit. Maecenas porttitor congue massa.'
    assert references == ''


def testGetDocxTextboxes(testClient):
    """
        Test if the getDOCXText function retrieves the text from the docx file without text from textboxes.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'textboxTest.docx')

    text, references = getDOCXText(fileDir)
    assert text == 'This is text outside a textbox.'
    assert references == ''


def testGetDocxEmptyFile(testClient):
    """
        Test if the getDOCXText function outputs an empty string when using an empty file.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'emptyFile.docx')

    text, references = getDOCXText(fileDir)
    assert text == ''
    assert references == ''


def testGetDocxCorruptedFile(testClient):
    """
        Test if the getDOCXText function outputs an empty string when using a corrupted file.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'corruptedFile.docx')

    text, references = getDOCXText(fileDir)
    assert text == ''
    assert references == ''


def testGetDocxInvalidFile(testClient):
    """
        Test if the getDOCXText function outputs an empty string when using a file path that does not exist.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'invalidFileName.docx')

    text, references = getDOCXText(fileDir)
    assert text == ''
    assert references == ''


def testGetDocxInvalidExtension(testClient):
    """
        Test if the getDOCXText function outputs an empty string when using a file that is not a docx file.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            text: String that contains the output of the getDOCXText function.
            references: String that contains the references extracted from the docx file.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'invalidFileExtension.pdf')

    text, references = getDOCXText(fileDir)
    assert text == ''
    assert references == ''


def testSubtractTextFromParagraphSingle(testClient):
    """
        Test if the subtractTextFromParagraph function outputs the correct text when reading a single paragraph with normal text.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            doc: Document that is read.
            para: Paragraph of the document.
            text: String that contains the output for normal text.
            references: String that contains the output for references text.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'paragraphtestsingle.docx')
    doc = docx.Document(fileDir)

    para = doc.paragraphs[0]  # First paragraph of document
    text, references = subtractTextFromParagraph(para, False, '', '')  # Run function
    assert text == 'Paragraph 1\n'
    assert references == ''


def testSubtractTextFromParagraphSingleRefs(testClient):
    """
        Test if the subtractTextFromParagraph function outputs the correct text when reading a single paragraph with references text.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            doc: Document that is read.
            para: Paragraph of the document.
            text: String that contains the output for normal text.
            references: String that contains the output for references text.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'paragraphtestsingle.docx')
    doc = docx.Document(fileDir)

    para = doc.paragraphs[0]  # First paragraph of document
    text, references = subtractTextFromParagraph(para, True, '', '')  # Run function
    assert text == ''
    assert references == 'Paragraph 1\n'


def testSubtractTextFromParagraphMultiple(testClient):
    """
        Test if the subtractTextFromParagraph function outputs the correct text when reading multiple paragraphs with normal text.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            doc: Document that is read.
            para1, para2: Paragraphs of the document.
            text1, text2: Strings that contain the output for normal text.
            references1, references2: Strings that contain the output for references text.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'paragraphtestmultiple.docx')
    doc = docx.Document(fileDir)

    para1 = doc.paragraphs[0]  # First paragraph of document
    para2 = doc.paragraphs[1]  # Second paragraph of document
    text1, references1 = subtractTextFromParagraph(para1, False, '', '')  # Run function
    assert text1 == 'Paragraph 1\n'
    assert references1 == ''

    text2, references2 = subtractTextFromParagraph(para2, False, '', '')  # Run function
    assert text2 == 'Paragraph 2.\n'
    assert references2 == ''


def testSubtractTextFromParagraphTextbox(testClient):
    """
        Test if the subtractTextFromParagraph function outputs the correct text when reading a paragraph with a textbox.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            doc: Document that is read.
            para: Paragraph of the document.
            text: String that contains the output for normal text.
            references: String that contains the output for references text.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'paragraphtesttextbox.docx')
    doc = docx.Document(fileDir)

    para = doc.paragraphs[0]  # First paragraph of document
    text, references = subtractTextFromParagraph(para, False, '', '')  # Run function
    assert text == '\n'
    assert references == ''


def testSubtractTextFromParagraphAppend(testClient):
    """
        Test if the subtractTextFromParagraph function outputs the correct text when appending
        text from a paragraph with normal text.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            doc: Document that is read.
            para: Paragraph of the document.
            text: String that contains the output for normal text.
            references: String that contains the output for references text.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'paragraphtestsingle.docx')
    doc = docx.Document(fileDir)

    para = doc.paragraphs[0]  # First paragraph of document
    text, references = subtractTextFromParagraph(para, False, 'Text\n', '')  # Run function
    assert text == 'Text\nParagraph 1\n'
    assert references == ''


def testSubtractTextFromParagraphAppendRefs(testClient):
    """
        Test if the subtractTextFromParagraph function outputs the correct text when appending
        text from a paragraph with references text.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            doc: Document that is read.
            para: Paragraph of the document.
            text: String that contains the output for normal text.
            references: String that contains the output for references text.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'paragraphtestsingle.docx')
    doc = docx.Document(fileDir)

    para = doc.paragraphs[0]  # First paragraph of document
    text, references = subtractTextFromParagraph(para, True, '', 'Text\n')  # Run function
    assert text == ''
    assert references == 'Text\nParagraph 1\n'


def testSubtractTextFromParagraphAppendBoth(testClient):
    """
        Test if the subtractTextFromParagraph function outputs the correct text when appending
        text from a paragraph with references text.
        Attributes:
            BASEDIR: Path of the folder where all files are stores
            fileDir: Path of the location where the file is stored.
            doc: Document that is read.
            para: Paragraph of the document.
            text: String that contains the output for normal text.
            references: String that contains the output for references text.
        Arguments:
            testClient:  The test client we test this for.
    """
    del testClient
    BASEDIR = os.path.abspath(os.path.dirname(__file__))
    fileDir = os.path.join(BASEDIR, 'paragraphtestmultiple.docx')
    doc = docx.Document(fileDir)

    para1 = doc.paragraphs[0]  # First paragraph of document
    para2 = doc.paragraphs[1]  # Second paragraph of document

    text1, references1 = subtractTextFromParagraph(para1, False, 'Text1\n', 'Text2\n')  # Run function
    text2, references2 = subtractTextFromParagraph(para2, True, text1, references1)  # Run function

    assert text2 == 'Text1\nParagraph 1\n'
    assert references2 == 'Text2\nParagraph 2.\n'
